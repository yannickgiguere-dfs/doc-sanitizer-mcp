"""Document text extraction for various file formats."""

import base64
import csv
import email
import io
from abc import ABC, abstractmethod
from dataclasses import dataclass
from email import policy
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class ExtractedDocument:
    """Result of document extraction."""
    content: str
    source_type: str
    metadata: dict


class ExtractionError(Exception):
    """Raised when document extraction fails."""
    pass


class BaseExtractor(ABC):
    """Base class for document extractors."""

    @property
    @abstractmethod
    def supported_extensions(self) -> list[str]:
        """List of supported file extensions (lowercase, with dot)."""
        pass

    @abstractmethod
    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        """Extract text from document content.

        Args:
            content: Raw bytes of the document
            filename: Original filename (used for metadata)

        Returns:
            ExtractedDocument with markdown-formatted content
        """
        pass


class PlainTextExtractor(BaseExtractor):
    """Extractor for plain text files."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".txt"]

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        return ExtractedDocument(
            content=text,
            source_type="text",
            metadata={"filename": filename},
        )


class WordExtractor(BaseExtractor):
    """Extractor for Microsoft Word documents."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            doc = DocxDocument(io.BytesIO(content))
        except Exception as e:
            raise ExtractionError(f"Failed to read Word document: {e}")

        lines = []

        for para in doc.paragraphs:
            # Handle headings
            if para.style.name.startswith("Heading"):
                level = 1
                try:
                    level = int(para.style.name.replace("Heading ", ""))
                except ValueError:
                    pass
                lines.append(f"{'#' * level} {para.text}")
            elif para.text.strip():
                lines.append(para.text)

        # Extract tables
        for table in doc.tables:
            lines.append("")  # Blank line before table
            lines.append(self._table_to_markdown(table))
            lines.append("")  # Blank line after table

        return ExtractedDocument(
            content="\n\n".join(lines),
            source_type="docx",
            metadata={"filename": filename},
        )

    def _table_to_markdown(self, table) -> str:
        """Convert a Word table to Markdown format."""
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

            # Add header separator after first row
            if i == 0:
                rows.append("|" + "|".join(["---"] * len(cells)) + "|")

        return "\n".join(rows)


class PDFExtractor(BaseExtractor):
    """Extractor for PDF documents."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".pdf"]

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            reader = PdfReader(io.BytesIO(content))
        except Exception as e:
            raise ExtractionError(f"Failed to read PDF document: {e}")

        pages = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"## Page {i}\n\n{text}")

        return ExtractedDocument(
            content="\n\n".join(pages),
            source_type="pdf",
            metadata={
                "filename": filename,
                "page_count": len(reader.pages),
            },
        )


class ExcelExtractor(BaseExtractor):
    """Extractor for Excel spreadsheets."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".xlsx", ".xls"]

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            excel_file = pd.ExcelFile(io.BytesIO(content))
        except Exception as e:
            raise ExtractionError(f"Failed to read Excel file: {e}")

        sheets = []
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            if not df.empty:
                markdown_table = self._dataframe_to_markdown(df)
                sheets.append(f"## Sheet: {sheet_name}\n\n{markdown_table}")

        return ExtractedDocument(
            content="\n\n".join(sheets),
            source_type="excel",
            metadata={
                "filename": filename,
                "sheet_count": len(excel_file.sheet_names),
                "sheets": excel_file.sheet_names,
            },
        )

    def _dataframe_to_markdown(self, df: pd.DataFrame) -> str:
        """Convert a pandas DataFrame to Markdown table format."""
        # Clean up column names
        headers = [str(col).replace("|", "\\|") for col in df.columns]

        lines = []
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---"] * len(headers)) + "|")

        for _, row in df.iterrows():
            cells = [str(val).replace("|", "\\|").replace("\n", " ") for val in row]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)


class CSVExtractor(BaseExtractor):
    """Extractor for CSV files."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".csv"]

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        try:
            df = pd.read_csv(io.StringIO(text))
        except Exception as e:
            raise ExtractionError(f"Failed to parse CSV file: {e}")

        markdown_table = self._dataframe_to_markdown(df)

        return ExtractedDocument(
            content=markdown_table,
            source_type="csv",
            metadata={
                "filename": filename,
                "row_count": len(df),
                "column_count": len(df.columns),
            },
        )

    def _dataframe_to_markdown(self, df: pd.DataFrame) -> str:
        """Convert a pandas DataFrame to Markdown table format."""
        headers = [str(col).replace("|", "\\|") for col in df.columns]

        lines = []
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---"] * len(headers)) + "|")

        for _, row in df.iterrows():
            cells = [str(val).replace("|", "\\|").replace("\n", " ") for val in row]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)


class EmailExtractor(BaseExtractor):
    """Extractor for email files (.eml)."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".eml"]

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            msg = email.message_from_bytes(content, policy=policy.default)
        except Exception as e:
            raise ExtractionError(f"Failed to parse email file: {e}")

        lines = []

        # Extract headers
        lines.append("## Email Headers\n")
        headers_to_extract = ["From", "To", "Cc", "Subject", "Date"]
        for header in headers_to_extract:
            value = msg.get(header, "")
            if value:
                lines.append(f"**{header}:** {value}")

        # Extract body
        lines.append("\n## Email Body\n")
        body = self._get_email_body(msg)
        lines.append(body)

        return ExtractedDocument(
            content="\n".join(lines),
            source_type="email",
            metadata={
                "filename": filename,
                "subject": msg.get("Subject", ""),
                "from": msg.get("From", ""),
                "date": msg.get("Date", ""),
            },
        )

    def _get_email_body(self, msg) -> str:
        """Extract the body text from an email message."""
        body = ""

        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        body = payload.decode("utf-8", errors="replace")
                        break
                elif content_type == "text/html" and not body:
                    # Fallback to HTML if no plain text
                    payload = part.get_payload(decode=True)
                    if payload:
                        body = payload.decode("utf-8", errors="replace")
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode("utf-8", errors="replace")

        return body


class DocumentExtractor:
    """Main extractor that routes to appropriate format-specific extractor."""

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    def __init__(self):
        self._extractors: dict[str, BaseExtractor] = {}
        self._register_extractors()

    def _register_extractors(self) -> None:
        """Register all available extractors."""
        extractors = [
            PlainTextExtractor(),
            WordExtractor(),
            PDFExtractor(),
            ExcelExtractor(),
            CSVExtractor(),
            EmailExtractor(),
        ]
        for extractor in extractors:
            for ext in extractor.supported_extensions:
                self._extractors[ext.lower()] = extractor

    def get_supported_extensions(self) -> list[str]:
        """Get list of supported file extensions."""
        return list(self._extractors.keys())

    def extract_from_base64(
        self,
        base64_content: str,
        filename: str
    ) -> ExtractedDocument:
        """Extract text from base64-encoded document content.

        Args:
            base64_content: Base64-encoded document bytes
            filename: Original filename (used to determine file type)

        Returns:
            ExtractedDocument with extracted content

        Raises:
            ExtractionError: If extraction fails
        """
        try:
            content = base64.b64decode(base64_content)
        except Exception as e:
            raise ExtractionError(f"Failed to decode base64 content: {e}")

        return self.extract(content, filename)

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        """Extract text from document content.

        Args:
            content: Raw document bytes
            filename: Original filename (used to determine file type)

        Returns:
            ExtractedDocument with extracted content

        Raises:
            ExtractionError: If extraction fails or file type is unsupported
        """
        # Check file size
        if len(content) > self.MAX_FILE_SIZE:
            raise ExtractionError(
                f"File exceeds maximum size of {self.MAX_FILE_SIZE // (1024*1024)}MB. "
                "Consider splitting the document into smaller parts."
            )

        # Get file extension
        ext = Path(filename).suffix.lower()

        if ext not in self._extractors:
            supported = ", ".join(self.get_supported_extensions())
            raise ExtractionError(
                f"Unsupported file type: {ext}. Supported types: {supported}"
            )

        extractor = self._extractors[ext]
        return extractor.extract(content, filename)

    def extract_from_file(self, file_path: str | Path) -> ExtractedDocument:
        """Extract text from a file on disk.

        Args:
            file_path: Path to the document file

        Returns:
            ExtractedDocument with extracted content
        """
        path = Path(file_path)
        if not path.exists():
            raise ExtractionError(f"File not found: {file_path}")

        content = path.read_bytes()
        return self.extract(content, path.name)
